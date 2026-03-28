import os
import random
import datetime
from docx import Document
from fpdf import FPDF

# 公司列表
companies = [
    "阿里巴巴集团", "腾讯控股", "百度集团", "京东集团", "华为技术有限公司",
    "小米集团", "字节跳动", "网易集团", "拼多多", "美团",
    "中国工商银行", "中国建设银行", "中国农业银行", "中国银行", "中国移动",
    "中国电信", "中国联通", "中国石油天然气集团", "中国石油化工集团", "中国建筑集团",
    "中国中铁", "中国铁建", "中国中车", "中国航天科技集团", "中国航天科工集团",
    "苹果公司", "微软公司", "谷歌公司", "亚马逊公司", "Meta",
    "特斯拉公司", "英特尔公司", "IBM", "甲骨文公司", "思科系统",
    "惠普公司", "戴尔科技", "通用电气", "福特汽车", "通用汽车",
    "波音公司", "辉瑞公司", "强生公司", "摩根大通", "高盛集团",
    "花旗集团", "瑞银集团", "德意志银行", "西门子", "宝马集团"
]

# 文档类型
document_types = [
    "销售合同", "采购合同", "付款凭证", "应付发票", "应收发票", "收款单"
]

# 生成随机金额
def generate_amount():
    return round(random.uniform(1000, 1000000), 2)

# 生成随机日期
def generate_date():
    start_date = datetime.date(2023, 1, 1)
    end_date = datetime.date(2024, 12, 31)
    days_between = (end_date - start_date).days
    random_days = random.randint(0, days_between)
    return start_date + datetime.timedelta(days=random_days)

# 生成随机合同编号
def generate_contract_id(doc_type):
    prefix = {
        "销售合同": "SC",
        "采购合同": "PC",
        "付款凭证": "PV",
        "应付发票": "PI",
        "应收发票": "RI",
        "收款单": "RS"
    }[doc_type]
    return f"{prefix}-{datetime.datetime.now().strftime('%Y%m%d')}-{random.randint(1000, 9999)}"

# 生成DOC文件
def generate_doc_file(file_path, doc_type, seller, buyer, amount, date, doc_id):
    doc = Document()
    
    # 添加标题
    doc.add_heading(doc_type, 0)
    
    # 添加基本信息
    doc.add_paragraph(f"文档编号: {doc_id}")
    doc.add_paragraph(f"日期: {date}")
    doc.add_paragraph()
    
    if doc_type in ["销售合同", "采购合同"]:
        # 合同内容
        doc.add_heading("合同双方", level=1)
        doc.add_paragraph(f"卖方: {seller}")
        doc.add_paragraph(f"买方: {buyer}")
        doc.add_paragraph()
        
        doc.add_heading("合同内容", level=1)
        doc.add_paragraph(f"{seller}与{buyer}就相关业务达成如下协议:")
        doc.add_paragraph(f"1. 交易金额: ￥{amount}")
        doc.add_paragraph(f"2. 交易日期: {date}")
        doc.add_paragraph(f"3. 支付方式: 银行转账")
        doc.add_paragraph(f"4. 交付方式: 电子交付")
        doc.add_paragraph()
        
        doc.add_heading("双方签字", level=1)
        doc.add_paragraph(f"{seller} (盖章)")
        doc.add_paragraph(f"{buyer} (盖章)")
    
    elif doc_type == "付款凭证":
        doc.add_heading("付款信息", level=1)
        doc.add_paragraph(f"付款方: {seller}")
        doc.add_paragraph(f"收款方: {buyer}")
        doc.add_paragraph(f"付款金额: ￥{amount}")
        doc.add_paragraph(f"付款日期: {date}")
        doc.add_paragraph(f"付款方式: 银行转账")
        doc.add_paragraph(f"转账凭证号: {random.randint(1000000000, 9999999999)}")
    
    elif doc_type == "应付发票":
        doc.add_heading("发票信息", level=1)
        doc.add_paragraph(f"开票方: {seller}")
        doc.add_paragraph(f"收票方: {buyer}")
        doc.add_paragraph(f"发票金额: ￥{amount}")
        doc.add_paragraph(f"开票日期: {date}")
        doc.add_paragraph(f"发票号码: {random.randint(10000000, 99999999)}")
        doc.add_paragraph(f"税率: 13%")
    
    elif doc_type == "应收发票":
        doc.add_heading("发票信息", level=1)
        doc.add_paragraph(f"开票方: {seller}")
        doc.add_paragraph(f"收票方: {buyer}")
        doc.add_paragraph(f"发票金额: ￥{amount}")
        doc.add_paragraph(f"开票日期: {date}")
        doc.add_paragraph(f"发票号码: {random.randint(10000000, 99999999)}")
        doc.add_paragraph(f"税率: 13%")
    
    elif doc_type == "收款单":
        doc.add_heading("收款信息", level=1)
        doc.add_paragraph(f"收款方: {seller}")
        doc.add_paragraph(f"付款方: {buyer}")
        doc.add_paragraph(f"收款金额: ￥{amount}")
        doc.add_paragraph(f"收款日期: {date}")
        doc.add_paragraph(f"收款方式: 银行转账")
        doc.add_paragraph(f"收款凭证号: {random.randint(1000000000, 9999999999)}")
    
    # 保存文件
    doc.save(file_path)

# 生成PDF文件
def generate_pdf_file(file_path, doc_type, seller, buyer, amount, date, doc_id):
    pdf = FPDF()
    pdf.add_page()
    
    # 设置字体
    pdf.set_font("Arial", "B", 16)
    
    # 添加标题
    pdf.cell(0, 10, doc_type, 0, 1, "C")
    pdf.ln(10)
    
    # 添加基本信息
    pdf.set_font("Arial", "", 12)
    pdf.cell(0, 10, f"文档编号: {doc_id}", 0, 1)
    pdf.cell(0, 10, f"日期: {date}", 0, 1)
    pdf.ln(10)
    
    if doc_type in ["销售合同", "采购合同"]:
        # 合同内容
        pdf.set_font("Arial", "B", 14)
        pdf.cell(0, 10, "合同双方", 0, 1)
        pdf.set_font("Arial", "", 12)
        pdf.cell(0, 10, f"卖方: {seller}", 0, 1)
        pdf.cell(0, 10, f"买方: {buyer}", 0, 1)
        pdf.ln(10)
        
        pdf.set_font("Arial", "B", 14)
        pdf.cell(0, 10, "合同内容", 0, 1)
        pdf.set_font("Arial", "", 12)
        pdf.cell(0, 10, f"{seller}与{buyer}就相关业务达成如下协议:", 0, 1)
        pdf.cell(0, 10, f"1. 交易金额: ￥{amount}", 0, 1)
        pdf.cell(0, 10, f"2. 交易日期: {date}", 0, 1)
        pdf.cell(0, 10, f"3. 支付方式: 银行转账", 0, 1)
        pdf.cell(0, 10, f"4. 交付方式: 电子交付", 0, 1)
        pdf.ln(10)
        
        pdf.set_font("Arial", "B", 14)
        pdf.cell(0, 10, "双方签字", 0, 1)
        pdf.set_font("Arial", "", 12)
        pdf.cell(0, 10, f"{seller} (盖章)", 0, 1)
        pdf.cell(0, 10, f"{buyer} (盖章)", 0, 1)
    
    elif doc_type == "付款凭证":
        pdf.set_font("Arial", "B", 14)
        pdf.cell(0, 10, "付款信息", 0, 1)
        pdf.set_font("Arial", "", 12)
        pdf.cell(0, 10, f"付款方: {seller}", 0, 1)
        pdf.cell(0, 10, f"收款方: {buyer}", 0, 1)
        pdf.cell(0, 10, f"付款金额: ￥{amount}", 0, 1)
        pdf.cell(0, 10, f"付款日期: {date}", 0, 1)
        pdf.cell(0, 10, f"付款方式: 银行转账", 0, 1)
        pdf.cell(0, 10, f"转账凭证号: {random.randint(1000000000, 9999999999)}", 0, 1)
    
    elif doc_type == "应付发票":
        pdf.set_font("Arial", "B", 14)
        pdf.cell(0, 10, "发票信息", 0, 1)
        pdf.set_font("Arial", "", 12)
        pdf.cell(0, 10, f"开票方: {seller}", 0, 1)
        pdf.cell(0, 10, f"收票方: {buyer}", 0, 1)
        pdf.cell(0, 10, f"发票金额: ￥{amount}", 0, 1)
        pdf.cell(0, 10, f"开票日期: {date}", 0, 1)
        pdf.cell(0, 10, f"发票号码: {random.randint(10000000, 99999999)}", 0, 1)
        pdf.cell(0, 10, f"税率: 13%", 0, 1)
    
    elif doc_type == "应收发票":
        pdf.set_font("Arial", "B", 14)
        pdf.cell(0, 10, "发票信息", 0, 1)
        pdf.set_font("Arial", "", 12)
        pdf.cell(0, 10, f"开票方: {seller}", 0, 1)
        pdf.cell(0, 10, f"收票方: {buyer}", 0, 1)
        pdf.cell(0, 10, f"发票金额: ￥{amount}", 0, 1)
        pdf.cell(0, 10, f"开票日期: {date}", 0, 1)
        pdf.cell(0, 10, f"发票号码: {random.randint(10000000, 99999999)}", 0, 1)
        pdf.cell(0, 10, f"税率: 13%", 0, 1)
    
    elif doc_type == "收款单":
        pdf.set_font("Arial", "B", 14)
        pdf.cell(0, 10, "收款信息", 0, 1)
        pdf.set_font("Arial", "", 12)
        pdf.cell(0, 10, f"收款方: {seller}", 0, 1)
        pdf.cell(0, 10, f"付款方: {buyer}", 0, 1)
        pdf.cell(0, 10, f"收款金额: ￥{amount}", 0, 1)
        pdf.cell(0, 10, f"收款日期: {date}", 0, 1)
        pdf.cell(0, 10, f"收款方式: 银行转账", 0, 1)
        pdf.cell(0, 10, f"收款凭证号: {random.randint(1000000000, 9999999999)}", 0, 1)
    
    # 保存文件
    pdf.output(file_path)

# 主函数
def main():
    output_dir = "D:\AI project\AI-search\测试文件目录"
    total_files = 1000
    
    # 确保输出目录存在
    os.makedirs(output_dir, exist_ok=True)
    print(f"输出目录: {output_dir}")
    print(f"开始生成{total_files}个商业交易文档...")
    
    try:
        for i in range(total_files):
            # 随机选择文档类型
            doc_type = random.choice(document_types)
            
            # 随机选择两家不同的公司
            seller, buyer = random.sample(companies, 2)
            
            # 生成随机数据
            amount = generate_amount()
            date = generate_date()
            doc_id = generate_contract_id(doc_type)
            
            # 生成文件名
            file_name = f"{doc_type}_{seller}_{buyer}_{i+1}"
            
            # 生成DOC文件
            doc_file_path = os.path.join(output_dir, f"{file_name}.docx")
            print(f"生成DOC文件: {doc_file_path}")
            generate_doc_file(doc_file_path, doc_type, seller, buyer, amount, date, doc_id)
            
            # 生成PDF文件
            pdf_file_path = os.path.join(output_dir, f"{file_name}.pdf")
            print(f"生成PDF文件: {pdf_file_path}")
            generate_pdf_file(pdf_file_path, doc_type, seller, buyer, amount, date, doc_id)
            
            if (i + 1) % 100 == 0:
                print(f"已生成{i + 1}个文件")
        
        print(f"生成完成！共生成{total_files}个DOC文件和{total_files}个PDF文件。")
    except Exception as e:
        print(f"错误: {str(e)}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    main()
